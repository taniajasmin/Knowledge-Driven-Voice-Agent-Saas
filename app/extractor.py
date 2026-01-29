# from pypdf import PdfReader
# from docx import Document

# def extract_text(file_path: str) -> str:
#     if file_path.endswith(".pdf"):
#         reader = PdfReader(file_path)
#         text = ""
#         for page in reader.pages:
#             text += page.extract_text() + "\n"
#         return text

#     elif file_path.endswith(".docx"):
#         doc = Document(file_path)
#         return "\n".join([para.text for para in doc.paragraphs])

#     else:
#         raise ValueError("Unsupported file type")

import os
from pypdf import PdfReader
from docx import Document

# Only needed for .doc conversion on Windows
import win32com.client


def convert_doc_to_docx(doc_path: str) -> str:
    """
    Uses Microsoft Word to convert old .doc file into .docx.
    This is the only reliable way to handle .doc files.
    """
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(os.path.abspath(doc_path))
    new_path = doc_path + "x"  # creates .docx

    doc.SaveAs(os.path.abspath(new_path), FileFormat=16)  # 16 = docx format
    doc.Close()
    word.Quit()

    return new_path


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts text from a valid .docx file.
    """
    document = Document(docx_path)
    return "\n".join([para.text for para in document.paragraphs])


def extract_text(file_path: str) -> str:
    """
    Main extractor that supports:
    - PDF
    - DOCX
    - DOC (auto converted to DOCX)
    """

    ext = os.path.splitext(file_path)[1].lower()

    # -------- PDF --------
    if ext == ".pdf":
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    # -------- DOCX --------
    if ext == ".docx":
        return extract_text_from_docx(file_path)

    # -------- DOC (convert first) --------
    if ext == ".doc":
        converted_path = convert_doc_to_docx(file_path)
        return extract_text_from_docx(converted_path)

    raise ValueError("Unsupported file type. Only PDF, DOC, DOCX are allowed.")
